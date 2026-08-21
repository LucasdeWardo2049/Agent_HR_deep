"""The complete deterministic orchestration behind the agent's single tool."""

import asyncio
import hashlib
import io
import json
import logging
from collections.abc import AsyncIterator, Awaitable, Callable
from contextlib import suppress
from dataclasses import dataclass, field
from functools import cache
from typing import Protocol
from urllib.parse import quote
from uuid import uuid4

import pymupdf
from pydantic import ValidationError
from agno.tools import tool
from agno.run.agent import CustomEvent
from docx import Document

from app.google_workspace import (
    DOCX_MIME,
    GOOGLE_DOC_MIME,
    PDF_MIME,
    GoogleWorkspaceClient,
    GoogleWorkspaceError,
    ReportArtifacts,
)
from app.llm import (
    GeminiPDFParser,
    LocalLLM,
    StructuredGenerator,
    StructuredOutputError,
    assess_candidate,
    normalize_candidate_profile,
    parse_candidate_text,
    parse_job_profile,
    professional_profile_payload,
)
from app.schemas import (
    CandidateAssessment,
    CandidateProfile,
    CriterionAssessment,
    JobProfile,
    ResumeFile,
    TalentSearchResult,
)
from app.settings import Settings, get_settings
from db.talent import TalentStore

logger = logging.getLogger(__name__)

# Bump this value whenever the structured professional profile changes. The
# version participates in the cache key so existing resumes are reprocessed
# once after a schema/prompt change, even when the Drive file is unchanged.
PROFILE_PIPELINE_VERSION = "professional-profile-v2"

# Bump these whenever the matching prompt or schema changes: the version is part
# of the cache key, so old entries miss instead of being served stale.
JOB_PROFILE_CACHE_VERSION = "job-profile-v1"
ASSESSMENT_CACHE_VERSION = "assessment-v1"


def _cache_key(version: str, payload: dict[str, object]) -> str:
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)
    return sha256_bytes(f"{version}\0{canonical}".encode())


def job_profile_cache_key(description: str) -> str:
    return _cache_key(JOB_PROFILE_CACHE_VERSION, {"description": " ".join(description.split())})


def assessment_cache_key(job_profile: JobProfile, candidate: CandidateProfile) -> str:
    """Key the assessment by the exact inputs the model would have received."""
    return _cache_key(
        ASSESSMENT_CACHE_VERSION,
        {
            "job_profile": job_profile.model_dump(mode="json"),
            "candidate": professional_profile_payload(candidate),
        },
    )


def is_sample_content(*texts: str | None) -> bool:
    """Sample answers must never be cached; they would poison later real runs."""
    from app.fixtures import SAMPLE_MARKER

    return any(SAMPLE_MARKER in (text or "") for text in texts)


class Store(Protocol):
    def get_existing_metadata(self, drive_file_ids: list[str]) -> dict[str, dict[str, str | None]]: ...

    def get_source_hash(self, drive_file_id: str) -> str | None: ...

    def get_candidate_id(self, drive_file_id: str) -> str | None: ...

    def upsert_profile(
        self,
        *,
        file_name: str,
        mime_type: str,
        source_hash: str,
        profile: CandidateProfile,
        parser_provider: str,
        model_name: str,
        fallback_used: bool,
    ) -> None: ...

    def update_source_metadata(
        self,
        *,
        drive_file_id: str,
        file_name: str,
        mime_type: str,
        drive_url: str | None,
    ) -> None: ...

    def list_profiles(self, drive_file_ids: set[str] | None = None) -> list[CandidateProfile]: ...

    def get_cached_json(self, cache_key: str) -> dict[str, object] | None: ...

    def put_cached_json(self, cache_key: str, kind: str, payload: dict[str, object]) -> None: ...

    def save_search(
        self,
        *,
        description: str,
        job_profile: JobProfile | None,
        assessments: list[CandidateAssessment],
        result: TalentSearchResult,
    ) -> None: ...


class Workspace(Protocol):
    async def list_resume_files(self) -> list[ResumeFile]: ...

    async def download_resume(self, file: ResumeFile) -> bytes: ...

    async def create_report(
        self,
        job_profile: JobProfile,
        assessments: list[CandidateAssessment],
        profiles: list[CandidateProfile],
    ) -> ReportArtifacts: ...


class PDFFallback(Protocol):
    async def parse(self, pdf_bytes: bytes) -> CandidateProfile: ...


@dataclass
class SyncStats:
    discovered: int = 0
    processed: int = 0
    skipped: int = 0
    failed: int = 0
    fallback_used: int = 0
    warnings: list[str] = field(default_factory=list)
    drive_file_ids: set[str] = field(default_factory=set)


ProgressCallback = Callable[[str, str, int | None, int | None], Awaitable[None]]


@dataclass
class TalentSearchProgressEvent(CustomEvent):
    """Progress safely streamed by AgentOS while the deterministic tool runs."""

    phase: str = ""
    label: str = ""
    current: int | None = None
    total: int | None = None

    def __str__(self) -> str:
        # Agno 2.8.5 concatenates generator item strings into the final tool
        # result. Progress belongs only to SSE; keep the final JSON untouched.
        return ""


async def _emit_progress(
    callback: ProgressCallback | None,
    phase: str,
    label: str,
    current: int | None = None,
    total: int | None = None,
) -> None:
    if callback is not None:
        await callback(phase, label, current, total)


def sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def profile_cache_hash(source: bytes | str) -> str:
    marker = source if isinstance(source, bytes) else f"modified:{source}".encode()
    payload = PROFILE_PIPELINE_VERSION.encode("utf-8") + b"\0" + marker
    return sha256_bytes(payload)


def extract_pdf_text(content: bytes) -> tuple[str, int]:
    with pymupdf.open(stream=content, filetype="pdf") as document:
        pages = [page.get_text("text") for page in document]
    return "\n".join(pages).strip(), len(pages)


def extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells if cell.text.strip()))
    return "\n".join(paragraphs).strip()


def is_extraction_usable(text: str, page_count: int = 1) -> bool:
    if page_count < 1:
        return False
    compact = "".join(character for character in text if not character.isspace())
    if len(compact) < max(200, page_count * 60):
        return False
    alphabetic_ratio = sum(character.isalpha() for character in compact) / len(compact)
    replacement_ratio = compact.count("\ufffd") / len(compact)
    printable_ratio = sum(character.isprintable() for character in compact) / len(compact)
    return alphabetic_ratio >= 0.45 and replacement_ratio <= 0.02 and printable_ratio >= 0.95


def normalize_assessment(
    raw: CandidateAssessment,
    candidate: CandidateProfile,
    job_profile: JobProfile,
) -> CandidateAssessment:
    by_id = {item.criterion_id: item for item in raw.criteria}
    criteria = [
        by_id.get(
            criterion.id,
            CriterionAssessment(
                criterion_id=criterion.id,
                status="unclear",
                notes="O modelo não retornou este critério; é necessária confirmação humana.",
            ),
        )
        for criterion in job_profile.criteria
    ]
    required_ids = {criterion.id for criterion in job_profile.criteria if criterion.required}
    required_supported = sum(item.criterion_id in required_ids and item.status == "supported" for item in criteria)
    required_total = len(required_ids)
    coverage = required_supported / required_total if required_total else 0.0
    return raw.model_copy(
        update={
            "candidate_id": candidate.candidate_id,
            "candidate_name": candidate.full_name,
            "criteria": criteria,
            "required_supported": required_supported,
            "required_total": required_total,
            "criteria_coverage": coverage,
        }
    )


class TalentService:
    def __init__(
        self,
        *,
        store: Store,
        workspace: Workspace,
        local_llm: StructuredGenerator,
        pdf_fallback: PDFFallback,
        settings: Settings | None = None,
    ) -> None:
        self.store = store
        self.workspace = workspace
        self.local_llm = local_llm
        self.pdf_fallback = pdf_fallback
        self.settings = settings or get_settings()

    async def _parse_job_profile_cached(self, description: str) -> JobProfile:
        cache_key = job_profile_cache_key(description)
        cached = await asyncio.to_thread(self.store.get_cached_json, cache_key)
        if cached is not None:
            try:
                return JobProfile.model_validate(cached)
            except ValidationError:
                # A row written by an older schema: fall through to the model.
                _log_event("job_profile_cache_invalid", status="miss")
        profile = await parse_job_profile(self.local_llm, description)
        if not is_sample_content(profile.title, profile.summary):
            await asyncio.to_thread(
                self.store.put_cached_json,
                cache_key,
                "job_profile",
                profile.model_dump(mode="json"),
            )
        return profile

    async def _read_cached_assessment(self, cache_key: str) -> CandidateAssessment | None:
        cached = await asyncio.to_thread(self.store.get_cached_json, cache_key)
        if cached is None:
            return None
        try:
            return CandidateAssessment.model_validate(cached)
        except ValidationError:
            _log_event("assessment_cache_invalid", status="miss")
            return None

    async def _write_cached_assessment(self, cache_key: str, assessment: CandidateAssessment) -> None:
        notes = [criterion.notes for criterion in assessment.criteria]
        if is_sample_content(assessment.professional_summary, *notes):
            return
        await asyncio.to_thread(
            self.store.put_cached_json,
            cache_key,
            "assessment",
            assessment.model_dump(mode="json"),
        )

    async def sync_profiles(self, progress: ProgressCallback | None = None) -> SyncStats:
        files = await self.workspace.list_resume_files()
        stats = SyncStats(
            discovered=len(files),
            drive_file_ids={file.drive_file_id for file in files},
        )
        await _emit_progress(progress, "syncing_resumes", "Sincronizando currículos", 0, len(files))
        drive_file_ids = [file.drive_file_id for file in files]
        existing_metadata = await asyncio.to_thread(self.store.get_existing_metadata, drive_file_ids)

        for index, file in enumerate(files, start=1):
            started = asyncio.get_running_loop().time()
            try:
                file_metadata = existing_metadata.get(file.drive_file_id, {})
                existing_hash = file_metadata.get("source_hash")
                source_hash = profile_cache_hash(file.modified_time) if file.modified_time else None

                if source_hash is not None and source_hash == existing_hash:
                    await asyncio.to_thread(
                        self.store.update_source_metadata,
                        drive_file_id=file.drive_file_id,
                        file_name=file.file_name,
                        mime_type=file.mime_type,
                        drive_url=file.drive_url,
                    )
                    stats.skipped += 1
                    continue
                content = await self.workspace.download_resume(file)
                if source_hash is None:
                    source_hash = profile_cache_hash(content)
                    if source_hash == existing_hash:
                        await asyncio.to_thread(
                            self.store.update_source_metadata,
                            drive_file_id=file.drive_file_id,
                            file_name=file.file_name,
                            mime_type=file.mime_type,
                            drive_url=file.drive_url,
                        )
                        stats.skipped += 1
                        continue
                profile, provider, fallback_used = await self._parse_resume(file, content)
                candidate_id = file_metadata.get("candidate_id")
                profile = normalize_candidate_profile(profile).model_copy(
                    update={
                        "candidate_id": candidate_id or f"candidate_{uuid4().hex}",
                        "source_drive_file_id": file.drive_file_id,
                        "source_drive_url": file.drive_url,
                    }
                )
                model_name = (
                    self.settings.gemini_pdf_model if fallback_used else self.settings.local_llm_model
                ) or "unknown"
                await asyncio.to_thread(
                    self.store.upsert_profile,
                    file_name=file.file_name,
                    mime_type=file.mime_type,
                    source_hash=source_hash,
                    profile=profile,
                    parser_provider=provider,
                    model_name=model_name,
                    fallback_used=fallback_used,
                )
                stats.processed += 1
                stats.fallback_used += int(fallback_used)
                _log_event(
                    "resume_parsed",
                    candidate_id=profile.candidate_id,
                    drive_file_id=file.drive_file_id,
                    parser_provider=provider,
                    model_name=model_name,
                    fallback_used=fallback_used,
                    duration_ms=round((asyncio.get_running_loop().time() - started) * 1000),
                    status="ok",
                )
            except Exception as exc:
                stats.failed += 1
                stats.warnings.append(f"{file.file_name}: {type(exc).__name__}")
                _log_event(
                    "resume_parse_failed",
                    drive_file_id=file.drive_file_id,
                    duration_ms=round((asyncio.get_running_loop().time() - started) * 1000),
                    status="failed",
                    error_type=type(exc).__name__,
                )
            finally:
                await _emit_progress(
                    progress,
                    "syncing_resumes",
                    f"Sincronizando currículos — {index} de {len(files)}",
                    index,
                    len(files),
                )
        return stats

    async def _parse_resume(self, file: ResumeFile, content: bytes) -> tuple[CandidateProfile, str, bool]:
        if file.mime_type == PDF_MIME:
            text, page_count = await asyncio.to_thread(extract_pdf_text, content)
            if is_extraction_usable(text, page_count):
                try:
                    return await parse_candidate_text(self.local_llm, text), "local_llm", False
                except StructuredOutputError:
                    pass
            return await self.pdf_fallback.parse(content), "gemini", True
        if file.mime_type == DOCX_MIME:
            text = await asyncio.to_thread(extract_docx_text, content)
        elif file.mime_type == GOOGLE_DOC_MIME:
            text = content.decode("utf-8", errors="replace").strip()
        else:
            raise ValueError(f"Unsupported resume MIME type: {file.mime_type}")
        if not is_extraction_usable(text):
            raise StructuredOutputError("The document does not contain enough usable text")
        return await parse_candidate_text(self.local_llm, text), "local_llm", False

    async def search(
        self,
        description: str,
        progress: ProgressCallback | None = None,
    ) -> TalentSearchResult:
        await _emit_progress(progress, "interpreting_job", "Interpretando os requisitos da vaga")
        job_profile = await self._parse_job_profile_cached(description)
        if not job_profile.is_actionable:
            await _emit_progress(progress, "needs_clarification", "Aguardando detalhes da vaga")
            return TalentSearchResult(
                status="needs_clarification",
                message=job_profile.clarification_question or "Descreva melhor os requisitos profissionais da vaga.",
            )

        search_id = f"search_{uuid4().hex}"
        assessments: list[CandidateAssessment] = []
        warnings: list[str] = []
        try:
            sync_stats = await self.sync_profiles(progress)
            warnings.extend(sync_stats.warnings)
            profiles = await asyncio.to_thread(self.store.list_profiles, sync_stats.drive_file_ids)
            semaphore = asyncio.Semaphore(self.settings.assessment_concurrency)

            async def assess(profile: CandidateProfile) -> tuple[CandidateAssessment, str | None]:
                cache_key = assessment_cache_key(job_profile, profile)
                cached = await self._read_cached_assessment(cache_key)
                if cached is not None:
                    return normalize_assessment(cached, profile, job_profile), None
                async with semaphore:
                    try:
                        raw = await assess_candidate(self.local_llm, job_profile, profile)
                        await self._write_cached_assessment(cache_key, raw)
                        return normalize_assessment(raw, profile, job_profile), None
                    except Exception as exc:
                        warning = f"{profile.full_name or profile.candidate_id}: {type(exc).__name__}"
                        fallback = CandidateAssessment(
                            candidate_id=profile.candidate_id,
                            candidate_name=profile.full_name,
                            criteria=[
                                CriterionAssessment(
                                    criterion_id=criterion.id,
                                    status="unclear",
                                    notes="Avaliação indisponível; é necessária revisão humana.",
                                )
                                for criterion in job_profile.criteria
                            ],
                            points_to_confirm=["Revisar manualmente os critérios profissionais desta pessoa."],
                        )
                        return normalize_assessment(fallback, profile, job_profile), warning

            total_profiles = len(profiles)
            await _emit_progress(
                progress,
                "assessing_candidates",
                "Analisando currículos",
                0,
                total_profiles,
            )
            tasks = [asyncio.create_task(assess(profile)) for profile in profiles]
            for completed_count, task in enumerate(asyncio.as_completed(tasks), start=1):
                assessment, warning = await task
                assessments.append(assessment)
                if warning:
                    warnings.append(warning)
                await _emit_progress(
                    progress,
                    "assessing_candidates",
                    f"Analisando currículos — {completed_count} de {total_profiles}",
                    completed_count,
                    total_profiles,
                )
            assessments.sort(key=lambda item: (item.candidate_name or "").casefold())
            if simulated_warning := _simulated_content_warning(job_profile, assessments):
                warnings.append(simulated_warning)
            await _emit_progress(progress, "generating_report", "Gerando planilhas e links para download")
            artifacts = await self.workspace.create_report(job_profile, assessments, profiles)
            excel_download_url = (
                f"{self.settings.public_app_url.rstrip('/')}/api/v1/talent/searches/{quote(search_id, safe='')}/xlsx"
            )
            result = TalentSearchResult(
                status="completed",
                message=(f"Busca concluída. {len(assessments)} perfis foram preparados para revisão humana."),
                search_id=search_id,
                candidates_analyzed=len(assessments),
                google_sheet_url=artifacts.google_sheet_url,
                excel_url=excel_download_url,
                excel_drive_file_id=artifacts.excel_file_id,
                warnings=warnings,
            )
        except Exception as exc:
            warning = str(exc) if isinstance(exc, GoogleWorkspaceError) else type(exc).__name__
            warnings.append(warning)
            result = TalentSearchResult(
                status="failed",
                message="Não foi possível concluir a busca. Verifique a configuração e tente novamente.",
                search_id=search_id,
                warnings=warnings,
            )
            _log_event("talent_search_failed", search_id=search_id, status="failed", error_type=type(exc).__name__)

        await asyncio.to_thread(
            self.store.save_search,
            description=description,
            job_profile=job_profile,
            assessments=assessments,
            result=result,
        )
        await _emit_progress(
            progress,
            "completed" if result.status == "completed" else "failed",
            "Busca concluída" if result.status == "completed" else "Não foi possível concluir a busca",
            result.candidates_analyzed,
            result.candidates_analyzed,
        )
        return result


@cache
def get_talent_service() -> TalentService:
    settings = get_settings()
    local_llm: StructuredGenerator = LocalLLM(settings)
    pdf_fallback: PDFFallback = GeminiPDFParser(settings)
    if settings.talent_sample_fallback:
        from app.fixtures import FallbackPDFParser, FallbackStructuredGenerator

        # The real model stays the first option; fixtures answer only after it
        # fails, and the resulting report is labelled as a demonstration sample.
        local_llm = FallbackStructuredGenerator(local_llm)
        pdf_fallback = FallbackPDFParser(pdf_fallback)
        _log_event("talent_sample_fallback_enabled", status="ok")
    return TalentService(
        store=TalentStore(),
        workspace=GoogleWorkspaceClient(settings),
        local_llm=local_llm,
        pdf_fallback=pdf_fallback,
        settings=settings,
    )


@tool(stop_after_tool_call=True)
async def search_talent_pool(description: str) -> AsyncIterator[TalentSearchProgressEvent | str]:
    """Busque evidências profissionais objetivas no banco de talentos configurado.

    Args:
        description: Descrição da vaga em linguagem natural fornecida pelo gestor.
    """
    events: asyncio.Queue[TalentSearchProgressEvent] = asyncio.Queue()

    async def publish(phase: str, label: str, current: int | None, total: int | None) -> None:
        await events.put(
            TalentSearchProgressEvent(
                phase=phase,
                label=label,
                current=current,
                total=total,
            )
        )

    search_task = asyncio.create_task(get_talent_service().search(description, progress=publish))
    event_task: asyncio.Task[TalentSearchProgressEvent] | None = None
    try:
        while not search_task.done():
            event_task = asyncio.create_task(events.get())
            done, _ = await asyncio.wait(
                (search_task, event_task),
                return_when=asyncio.FIRST_COMPLETED,
            )
            if event_task in done:
                yield event_task.result()
            else:
                event_task.cancel()
                with suppress(asyncio.CancelledError):
                    await event_task
            event_task = None

        while not events.empty():
            yield events.get_nowait()
        result = await search_task
        yield result.model_dump_json()
    finally:
        if event_task is not None and not event_task.done():
            event_task.cancel()
        if not search_task.done():
            search_task.cancel()
            with suppress(asyncio.CancelledError):
                await search_task


def _simulated_content_warning(
    job_profile: JobProfile,
    assessments: list[CandidateAssessment],
) -> str | None:
    """Label fixture-sourced content so a sample report is never taken for a real one."""
    from app.fixtures import SAMPLE_MARKER

    texts = [job_profile.title, *(assessment.professional_summary or "" for assessment in assessments)]
    if not any(SAMPLE_MARKER in text for text in texts):
        return None
    return "Resultado gerado com dados de amostra."


def _log_event(event: str, **fields: object) -> None:
    logger.info(json.dumps({"event": event, **fields}, ensure_ascii=False, default=str))
