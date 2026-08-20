"""Opt-in real integration test.

Run only in an isolated account/window:
RUN_TALENT_E2E=1 pytest -m e2e tests/test_e2e_external.py
"""

import io
import os
from pathlib import Path
from urllib.parse import parse_qs, urlparse
from uuid import uuid4

import pymupdf
import pytest
from docx import Document
from openpyxl import load_workbook

from app.google_workspace import GoogleWorkspaceClient
from app.llm import GeminiPDFParser, LocalLLM
from app.settings import get_settings
from app.talent import TalentService
from db import TalentStore, init_talent_tables

pytestmark = [
    pytest.mark.e2e,
    pytest.mark.skipif(os.getenv("RUN_TALENT_E2E") != "1", reason="set RUN_TALENT_E2E=1"),
]


def _drive_id(url: str) -> str:
    parsed = urlparse(url)
    query_id = parse_qs(parsed.query).get("id")
    if query_id:
        return query_id[0]
    parts = [part for part in parsed.path.split("/") if part]
    if "d" in parts:
        return parts[parts.index("d") + 1]
    raise AssertionError(f"Could not extract a Drive id from {url}")


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Marina Synthetic", 0)
    document.add_paragraph(
        "Professional software engineer with five years of Python, FastAPI, PostgreSQL, "
        "automated testing, APIs, Docker, and collaborative product delivery experience. " * 3
    )
    document.save(str(path))


def _create_image_only_pdf(path: Path) -> None:
    source = pymupdf.open()
    page = source.new_page(width=900, height=1200)
    page.insert_textbox(
        pymupdf.Rect(70, 70, 830, 1100),
        (
            "Rafael Synthetic\n\nData analyst with SQL, Python and dashboard experience. "
            "Built business reports and validated data quality for three years. " * 5
        ),
        fontsize=22,
    )
    image = page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5)).tobytes("png")
    source.close()

    scanned = pymupdf.open()
    scanned_page = scanned.new_page(width=900, height=1200)
    scanned_page.insert_image(scanned_page.rect, stream=image)
    scanned.save(path)
    scanned.close()


async def test_real_drive_sheets_gemini_and_xlsx(tmp_path: Path) -> None:
    settings = get_settings()
    settings.require_google()
    assert settings.gemini_api_key and settings.gemini_pdf_model
    init_talent_tables()

    bootstrap = GoogleWorkspaceClient(settings)
    created_ids: list[str] = []
    folder_id: str | None = None
    try:
        suffix = uuid4().hex[:10]
        folder_id = await bootstrap.create_folder(
            f"talent-e2e-{suffix}",
            settings.google_drive_talent_pool_folder_id,
        )
        docx_path = tmp_path / f"marina-{suffix}.docx"
        pdf_path = tmp_path / f"rafael-{suffix}.pdf"
        _create_docx(docx_path)
        _create_image_only_pdf(pdf_path)
        created_ids.extend(
            [
                await bootstrap.upload_file(docx_path, folder_id),
                await bootstrap.upload_file(pdf_path, folder_id),
            ]
        )

        isolated_settings = settings.model_copy(
            update={
                "google_drive_talent_pool_folder_id": folder_id,
                "results_drive_folder_id": folder_id,
            }
        )
        workspace = GoogleWorkspaceClient(isolated_settings)
        service = TalentService(
            store=TalentStore(),
            workspace=workspace,
            local_llm=LocalLLM(isolated_settings),
            pdf_fallback=GeminiPDFParser(isolated_settings),
            settings=isolated_settings,
        )

        result = await service.search(
            "Pessoa profissional de dados ou backend com Python obrigatório; SQL ou FastAPI desejável."
        )

        assert result.status == "completed", result.model_dump()
        assert result.candidates_analyzed == 2
        assert result.google_sheet_url and result.excel_url and result.excel_drive_file_id
        sheet_id = _drive_id(result.google_sheet_url)
        excel_id = result.excel_drive_file_id
        created_ids.extend([sheet_id, excel_id])
        assert await workspace.get_sheet_names(sheet_id) == ["Summary", "Criteria", "Candidates"]
        excel_bytes = await workspace.download_file(excel_id)
        workbook = load_workbook(io.BytesIO(excel_bytes), read_only=True)
        assert workbook.sheetnames == ["Summary", "Criteria", "Candidates"]
    finally:
        for file_id in reversed(created_ids):
            try:
                await bootstrap.delete_file(file_id)
            except Exception:
                pass
        if folder_id:
            try:
                await bootstrap.delete_file(folder_id)
            except Exception:
                pass
